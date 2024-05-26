import time
t1 = time.time()

import pdfplumber
import re
import io
from pykospacing import Spacing
from docx import Document
from docx.shared import Inches
from typing import List, Union
from docx.oxml.ns import qn
from kiwipiepy import Kiwi

PAGE_ENDED = '@'*10

ACCESS_PTRN = re.compile(r"[^가-힣ㄱ-ㅎ0-9a-zA-Z\s@_!#$%^&*()-=+<>?/\|}{~:]")
# REPLACE_PTRN = 

def replace_to_single_double_quote(text: str) -> str:
    """
    (와 )도 SSO, SSC로 분류되는데, 이건 사람이 직접 다 확인하고, 제거해야 한다.
    특수문자 보관함:
    ●
    """
    # 왜 안됨?
    text = text.replace("\u201C", "'")\
               .replace("  ", " ")\
               .replace("\u201D", "'")\
               .replace("\u2018", "'")\
               .replace("\u2019", "'")\
               .replace("\u201D", "'")\
               .replace('「', "'")\
               .replace('」', "'")\
               .replace('『', "'")\
               .replace("』", "'")\
               .replace("〈", "'")\
               .replace("〉","'") \
               .replace("・ ", ", ") \
               .replace(" ・", ", ") \
               .replace("・", ", ") \
               .replace("～", "~")\
               .replace("＇", "'")\
               .replace("ʼ", "'")\
    
    return text
               

def only_acessed_char(text:str) -> str:
    return ACCESS_PTRN.sub("", text).strip()


def need_newline(word):
    # lsquot = '\u2018'
    # rsquot = '\u2019'
    # ldquot = '\u201C'
    # rdquot = '\u201D'
    # "— "는 파이썬에서 "\u2014 "
    if re.search(r'[.!?][\u201D\u2019]', word):
        return True
    else:
        return False


def convert_fontcolor_to_tuple(fontcolor: any) -> tuple: # Union[tuple(int, ...), List[int]]
    # black
    if fontcolor == (0,0,0,1):
        return (1,0,0)
    
    # white
    if fontcolor == (0,0,0,0):
        return (0,1,0)

    # others
    return (0,0,1)


class ScriptExtractor:
    '''
    #### class variable: share all things inside the ScriptExtractor class for a book. 
    document: to handle .docx extension.
    spacing: whether is neccessary for space or not, checking the space
    must be needed between the first word of newline 
    and last word of prev line.
    title: title of book.
    categorize_by_fontsize: classify the words whether insert newline or not.
    contents: contain contents in this book.
    layout_page: set standard for determining the coordinates of page.
    watermark: if there is watermark, it should be eliminated.
    remove_english_footnote: remove english footnote for korean book.
    remove_english_intext_citation: remove english intext citation for korean book.
    ------
    '''    
    path = ''
    title = ''
    contents = [] # after works.
    layout_pages = []
    categorize_by_font = {}
    remove_english_footnote = True
    remove_english_intext_citation = True
    watermark = False
    output_path = ""
    even_layout_location = {'leftmost': 10000, 
                        'topmost':10000, 
                        'rightmost':0, 
                        'bottommost':0,
                        }
    odd_layout_location = {'leftmost': 10000, 
                        'topmost':10000, 
                        'rightmost':0, 
                        'bottommost':0,
                        }
    layout_location = {'leftmost': None, 
                        'topmost':None, 
                        'rightmost':None, 
                        'bottommost':None,
                        }
    kiwi = Kiwi(model_type='sbg')
    line_height_threshold = 15
    excluded_lines_count = 1
    NEWLINE_SPACE = 5 # Could be Hypertuned
    PARAGRAPH_SPACE = 5 # Could be Hypertuned
    LINE_HEIGHT_THRESHOLD = 25 # Could be Hypertuned
    
    @property
    def get_layout_location(self):
         return self.__class__.layout_location
    
    @property
    def get_categorize_by_font(self):
         return self.__class__.categorize_by_font

    @classmethod
    def set_option(cls, path:str, output_path:str, title:str, layout_page:int, LINE_HEIGHT_THRESHOLD=15,
                    PARAGRAPH_SPACE = 5,
                    excluded_lines_count=1, watermark=False, 
                    remove_english_footnote=True,
                    remove_english_intext_citation=True) -> None:
        
        cls.output_path = output_path
        cls.path = path
        cls.title = title
        cls.layout_pages = [layout_page, layout_page+1]
        cls.watermark = watermark
        cls.remove_english_footnote = remove_english_footnote
        cls.remove_english_intext_citation = remove_english_intext_citation
        cls.excluded_lines_count = excluded_lines_count
        cls.LINE_HEIGHT_THRESHOLD = LINE_HEIGHT_THRESHOLD
        cls.PARAGRAPH_SPACE = PARAGRAPH_SPACE
        # box should be considered.

        # later, it may changes.
        #         
        custom_layout_by_font_list = []
        user_font_input = str(input("더 입력하시겠습니까? 그러시려면 전과 같이 입력해주세요. \n \
                                        아니라면 '끝'이라고 입력해주세요: "))
        while user_font_input != "끝":
            user_font = user_font_input.replace(' ', '').split(',')
            user_font[1] = round(float(user_font[1]),1)
           
            if user_font[2] == '검':
                user_font[2] = (1,0,0)
            elif user_font[2] == '흰':
                user_font[2] = (0,1,0)
            else:
                user_font[2] = (0,0,1)
            
            custom_layout_by_font_list.append(user_font)
            user_font_input = str(input("더 입력하시겠습니까? 그러시려면 전과 같이 입력해주세요. \n \
                                        아니라면 '끝'이라고 입력해주세요."))
            
        # custom_layout_by_font_list: [['본문', 11, (0,0,1)], ]
        cls.categorize_by_font = {(i[1],i[2]):i[0] for i in custom_layout_by_font_list}
        # {"11, (1,0,0)": "본문"}


    def __init__(self, pages:List[int]) -> None:
        self.pages = pages
        self.words_in_range = [] # make word file by instance variable, not accumulating.    
        self.last_word = {
                        'x0': None, 
                        'top':None, 
                        'x1':None, 
                        'bottom':None,
                        'text':''
                        }
        self.prev_line = {
                        'x0': None, 
                        'top':None, 
                        'x1':None, 
                        'bottom':None,
                        'text':''
                        }
        self.image_page = []
        self.images = {}

    @classmethod
    # Input: 모든 줄글.
    # Output: 모든 줄글에 레이블링한 것이 결과가 된다.
    # make model로 모델을 만들고, 이걸 cls.method나 variable로 저장.
    # 이후, 이걸 self.page_range에 맞게 한번에 대입하도록 만듦.
    def classify_word_by_font(
        cls, 
        word:str, 
        fontsize: float,
        x0: float,
        top: float,
        fontcolor: any # tuple(int,...) | List[int]
        ) -> str: # tuple(int) | List[int]

        # 서문은 어떻게...?
        pages = range(int(input("목차를 제외하고 페이지 시작을 입력해주세요.")), int(input("끝 페이지를 입력해주세요"))+1)
        with pdfplumber.open(cls.path, pages=pages, laparams=None) as pdf:  
            page_shifted = False
            
            leftmost, topmost, rightmost, bottommost = [i for i in cls.layout_location.values()]



            for idx, page in enumerate(pdf.pages):
                page = page.crop((leftmost, topmost, rightmost, bottommost))             

                for word in page.extract_words(use_text_flow=self.__class__.watermark,
                                                extra_attrs=['size', 'fontname', 'stroking_color', 'non_stroking_color']):
                    pass

    @classmethod
    def set_layout(cls):
        with pdfplumber.open(cls.path, pages=cls.layout_pages, laparams=None) as pdf:   
            cls.layout_location['leftmost'] = 100000 # max position.
            cls.layout_location['topmost'] = 100000 # max position.
            cls.layout_location['rightmost'] = 0
            cls.layout_location['bottommost'] = 0

            last_word_size = 0
            for page in pdf.pages:
                for word in page.extract_words():
                    cls.layout_location['leftmost'] = min(cls.layout_location['leftmost'], word['x0']) 
                    cls.layout_location['topmost'] = min(cls.layout_location['topmost'], word['top']) 
                    cls.layout_location['rightmost'] = max(cls.layout_location['rightmost'], word['x1']) 
                    cls.layout_location['bottommost'] = max(cls.layout_location['bottommost'], word['bottom'])
                last_word_size = page.chars[-1]['size'] # page number size.
            
            cls.layout_location['topmost'] -= 1.5*last_word_size
            cls.layout_location['bottommost'] -= 1.5*last_word_size

    
    def set_new_layout(cls, fontsize_times={'top': 1.5, 'bottom': 1.5}):
        """
        Set the page layout by odd and even
        without page_num or upper decoration.
        """
        with pdfplumber.open(cls.path, pages=cls.layout_pages, laparams=None) as pdf:   
            # 마지막 line이 몇 개인지 사용자 입력, 테스트 등으로 파악해두는게 필요할듯.
            last_word_size = 0
            for page in pdf.pages:
                page_lines = page.extract_text_lines(layout=True)
                if page.page_number % 2 == 0:
                    for idx, line_dict in enumerate(page_lines):
                        if len(page_lines) - cls.excluded_lines_count <= idx:
                            break
                        # 마지막 line이 몇 개인지
                        cls.even_layout_location['leftmost'] = min(cls.even_layout_location['leftmost'], line_dict['x0'])
                        cls.even_layout_location['topmost'] = min(cls.even_layout_location['topmost'], line_dict['top']) 
                        cls.even_layout_location['rightmost'] = max(cls.even_layout_location['rightmost'], line_dict['x1']) 
                        cls.even_layout_location['bottommost'] = max(cls.even_layout_location['bottommost'], line_dict['bottom'])
                else:
                    for idx, line_dict in enumerate(page_lines):
                        if len(page_lines) - cls.excluded_lines_count <= idx:
                            break
                        cls.odd_layout_location['leftmost'] = min(cls.odd_layout_location['leftmost'], line_dict['x0'])
                        cls.odd_layout_location['topmost'] = min(cls.odd_layout_location['topmost'], line_dict['top']) 
                        cls.odd_layout_location['rightmost'] = max(cls.odd_layout_location['rightmost'], line_dict['x1']) 
                        cls.odd_layout_location['bottommost'] = max(cls.odd_layout_location['bottommost'], line_dict['bottom'])
                
                # last_word_size = page.chars[-1]['size'] # page number size.
            
            # cls.even_layout_location['topmost'] += fontsize_times['top']*last_word_size
            # cls.even_layout_location['bottommost'] -= fontsize_times['bottom']*last_word_size
            # cls.odd_layout_location['topmost'] += fontsize_times['top']*last_word_size
            # cls.odd_layout_location['bottommost'] -= fontsize_times['bottom']*last_word_size

    
    def add_newline_by_quotes(self):
        newline_pattern = r'[.!?][\u201D\u2019]' 
        fake_pattern = r'[\u201D\u2019]'
        quotes_indices = []
        for idx, word in enumerate(self.words_in_range):
            if '\u2018' in word or '\u201C' in word:
                # 어차피 무조건 [0]에 위치하기에 그냥 idx만 append.
                sum_count = word.count('\u2018')+word.count('\u201C')
                [quotes_indices.append(idx) for i in range(sum_count)]
            # 만약 match되면, 그리고 quotes_indices이면
            if re.search(fake_pattern, word):
                if re.search(newline_pattern, word) and quotes_indices:
                        # 해당하는 곳 앞에 '\n'를 넣어줌.
                        self.words_in_range[quotes_indices[0]] = '\n' + self.words_in_range[quotes_indices[0]]
                        # 맨 처음 들어온 것을 지움. queue로도 구현 가능할듯.
                        quotes_indices.pop(0)
                else:
                        # 그냥 강조 부분이라면 삭제함. 뒤에서부터 삭제해야 함.
                        quotes_indices.pop()


    def spacing_condition(self, x0: float) -> bool:
        # to be more specifically, divided by page to check spacing.
        if self.last_word['x1'] is None:
            return False
        
        if self.last_word['x1'] - x0 < 190 : # should be hypertuned.
            return False
        
        if re.search(r'[.!?]', self.last_word['text']) is not None:
            return False
        
        return True
    
    def over_line_height(self, top:float) -> bool: # line_height 초과 조건.
        if self.last_word['bottom'] is None:
            return False
        if top - self.last_word['bottom'] <= self.__class__.line_height_threshold:
            return False
        
        return True

    def update_last_word(self, word):
            for key in self.last_word.keys():
                        self.last_word[key] = word[key]


    def update_prev_line(self, word):
            for key in self.prev_line.keys():
                        self.prev_line[key] = word[key]

    def eliminate_watermark(self):
        for i in range(len(self.__class__.watermark)):
            if self.__class__.watermark[0] == self.words_in_range[-1][-1]:
                self.words_in_range[-1] = self.words_in_range[-1][0:-1]
                if self.words_in_range[-1][0] == '*':
                    self.words_in_range[-1] = self.words_in_range[-1][1:]
            else:
                self.words_in_range.pop()

    def seperate_sentence(self):
        n_pattern = r'[.!?][\u201D\u2019]'
        f_pattern = r'[\u201D\u2019]'
        quotes_indices = []
        for idx, word in enumerate(self.words_in_range):
            if '\u2018' in word or '\u201C' in word:
                # 어차피 무조건 [0]에 위치하기에 그냥 idx만 append.
                sum_count = word.count('\u2018')+word.count('\u201C')
                quotes_indices = [idx for i in range(sum_count)]
            # 만약 match되면, 그리고 quotes_indices이면
            if re.search(f_pattern, word):
                if re.search(n_pattern, word) and quotes_indices:
                    # 해당하는 곳 앞에 '\n'를 넣어줌.
                    self.words_in_range[quotes_indices[0]] = '\n' + self.words_in_range[quotes_indices[0]]
                    # 맨 처음 들어온 것을 지움. queue로도 구현 가능할듯.
                    quotes_indices.pop(0)
                else:
                    # 그냥 강조 부분이라면 삭제함. 뒤에서부터 삭제해야 함.
                    quotes_indices.pop()

    
    def read_pdf(self, double_quote_to_single_quote=False):
        with pdfplumber.open(self.__class__.path, pages=self.pages, laparams=None) as pdf:  
            page_shifted = False
            spacing = Spacing()
            # pre: set_layout() 이전에 되어 있어야 함.

            font_prop = {}
            
            leftmost, topmost, rightmost, bottommost = [i for i in self.__class__.layout_location.values()]
            for idx, page in enumerate(pdf.pages):
                page = page.crop((leftmost, topmost, rightmost, bottommost))
                
                # if page.lines:
                #     print('line이 있는 페이지',page.page_number)
                # if page.rects:
                #     print('rects이 있는 페이지',page.page_number)
                #     for rect in page.rects:
                #         print(rect['x0'], rect['top'], rect['x1'], rect['bottom'])
                # if page.curves:
                #     print('curves이 있는 페이지',page.page_number)
                #     for curve in page.curves:
                #         print(curve['x0'], curve['y0'], curve['x1'], curve['y1'])
                # if page.images:
                #     print('images이 있는 페이지',page.page_number)
                # if page.annots:
                #     print('annots가 있는 페이지',page.page_number)
                # # print(page)
                

                

                # check table instance
                if (len(page.lines) >= 4 
                    or len(page.curves) >= 4 
                    or page.images
                    ):
                    self.image_page.append(str(page.page_number))
                    self.images[idx] = page.to_image(resolution=128) # later, save the image to .word file.

                for word in page.extract_words(use_text_flow=self.__class__.watermark,
                                               x_tolerance=5,
                                               y_tolerance=5,
                                            #  split_at_punctutation=True,
                                            extra_attrs=['size', 'stroking_color', 'non_stroking_color'],
                                            ):
                    
                    # print(f"word: {word}")

                    # key = (round(word['size'],1), convert_fontcolor_to_tuple(word['non_stroking_color']))
                    # if key not in font_prop:
                    #     font_prop[key] = {'count':1, 'page_num':page.page_number, 'text':word['text']}
                    # else:
                    #     font_prop[key]['count'] += 1
                    # new_dict = dict(sorted(font_prop.items(), key=lambda x: x[1]['count'], reverse=True))
                    # for key, value in new_dict.items():
                    #     print(key, value)

                    x0, top, x1, bottom = map(lambda x : round(x, 2), (word['x0'], word['top'], word['x1'], word['bottom']))
                    
                    cur_word = word['text']

                    if double_quote_to_single_quote:
                        cur_word = replace_to_single_double_quote(word['text'])                                
                    
                    # 완전히 챕터가 종료되는 경우. 이건 나중에...
                    if page_shifted is True:
                        page_shifted = False
                        if (re.search(r'[.!?]', self.last_word['text']) # .!?가 있었고, 
                            and self.last_word['bottom'] < self.__class__.layout_location['bottommost'] - 55 # bottom에 놓는 것이 
                            ):
                            self.words_in_range.append('\n' + cur_word)
                            self.update_last_word(word)
                            continue
                        
                        if (re.search(r'[.!?]', self.last_word['text']) 
                            and self.last_word['x1'] < self.__class__.layout_location['rightmost'] - 10
                            ):                               
                            self.words_in_range.append('\n'+cur_word)
                            self.update_last_word(word)
                            continue
                        
                        if self.spacing_condition(x0):
                            self.words_in_range.pop()                               
                            cur_word = spacing(self.words_in_range.pop() + cur_word)
                            # append PAGE_ENDED in front of the spaced word.
                            self.words_in_range.append(PAGE_ENDED)
                            self.words_in_range.append(cur_word)
                            self.update_last_word(word)
                            continue
                    
                            
                    if self.spacing_condition(x0) and (not self.over_line_height(top)):
                        # No PAGE_ENEDED.
                        cur_word = spacing(self.words_in_range.pop() + cur_word)

                    if self.over_line_height(top): # 이것도 개인 설정.
                        cur_word = '\n'+ cur_word
                    
                    if need_newline(cur_word): 
                        cur_word = cur_word+'\n'
                    
                    self.words_in_range.append(cur_word)
                    
                    # self.last_word update.
                    self.update_last_word(word)
                
                # 페이지 전환.
                page_shifted = True
                    
                # 워터마크 제거. 보완필요.
                if self.__class__.watermark:
                    self.eliminate_watermark()

                self.words_in_range.append(PAGE_ENDED)
        
        print("pdf reading done.")
        # self.seperate_sentence() # 이거 녹여내면 속도 급상승할텐데...


    def read_pdf_v2(self):
        #TODO: sample 페이지로 layout 설정하는 과정 무조건 필요!
        # 홀수, 짝수 페이지로 layout 설정하고,
        # 그리고 나서 line별로 다 replace로 바꾸고,
        # line 별 x와 y있고, 맨 마지막이 .!?이라면 여기까지 glue 써주고 txt에 써주고.
        # 만약 line 별 x가 연속적으로 최대값 전 즈음이면 그냥 붙여준 다음 .붙여주기도 하고.
        # (보통 소주제 같은 것들)
        # line에서 x가 맨 끝이 아닌데 .이 있으면 또 . 넣어주고.
        # 만약 line에서 y가 큰 간격이라면 사이에 엔터 넣어주고.
        # 만약 문단 변화가 있는 곳에서 엔터를 넣는 옵션을 준다면 거기서 엔터를 또 넣어주고.
        with pdfplumber.open(self.__class__.path, pages=self.pages, laparams=None) as pdf, \
            open(self.__class__.output_path + self.__class__.title+ '_' + str(self.pages[0]) +'-' + str(self.pages[-1])+'.txt', 'w', encoding='utf-8') as ofile \
            :  
            # pre: set_new_layout() 이전에 되어 있어야 함.


            # enter와 .이 잘못 들어가 있음.
            # 또 1장과 같이 다른 글씨체면 문제가 생긴다
            # 마지막으로 작은 영어 글씨체면 맨 마지막에 간다.
            font_prop = {} 
            leftmost, topmost, rightmost, bottommost = (0,0,0,0)
            lines = []
            

            for idx, page in enumerate(pdf.pages):
                # 홀수, 짝수 페이지 별로 layout 설정.
                leftmost, topmost, rightmost, bottommost = (0, 0, 0, 0)
                if page.page_number % 2 == 0:
                    leftmost, topmost, rightmost, bottommost = [i for i in self.__class__.even_layout_location.values()]
                else:
                    leftmost, topmost, rightmost, bottommost = [i for i in self.__class__.odd_layout_location.values()]

                page = page.crop((leftmost, topmost, rightmost, bottommost))        
                
                # if page.lines:
                #     print('line이 있는 페이지',page.page_number)
                # if page.rects:
                #     print('rects이 있는 페이지',page.page_number)
                #     for rect in page.rects:
                #         print(rect['x0'], rect['top'], rect['x1'], rect['bottom'])
                # if page.curves:
                #     print('curves이 있는 페이지',page.page_number)
                #     for curve in page.curves:
                #         print(curve['x0'], curve['y0'], curve['x1'], curve['y1'])
                # if page.images:
                #     print('images이 있는 페이지',page.page_number)
                # if page.annots:
                #     print('annots가 있는 페이지',page.page_number)
                # # print(page)
                
                # check table, image, etc ... instance
                if (len(page.lines) >= 4 
                    or len(page.curves) >= 4 
                    or page.images
                    ):
                    self.image_page.append(str(page.page_number))
                    self.images[idx] = page.to_image(resolution=128) # later, save the image to .word file.

                prev_line = None
                # page_lines = page.extract_text_lines(layout=True, strip=True)
                page_lines = page.extract_text_lines(layout=False)
                for idx, line_dict in enumerate(page_lines):
                    line_text = replace_to_single_double_quote(line_dict['text']) # 그리고 나서 line별로 다 replace로 바꾸고,                    #TODO: 숫자가 있고, 뒤에 글자(한글, 영어, 특수문자 존재 시 붙여쓰기 만들기)
                    line_text = only_acessed_char(line_text)
                    line_dict['text'] = line_text
                    file_wrote_flag = False

                    if not line_text:
                        continue
                    
                    if "누가" in line_text and "끌고" in line_text and "가고" in line_text:
                        print(line_text)
                        



                    lines.append(line_text)
                    sent = ""
                    #TODO: 따옴표까지도 따로 분리하는 옵션은 이후 추가 예정.

                    if (line_text[-1] in ".!?"
                        or line_text[-2] in ".!?"
                        ):
                        sent = self.__class__.kiwi.glue(lines)                        
                        file_wrote = True
                    # elif self.prev_line['x1'] < rightmost - NEWLINE_SPACE and :
                        
                    ## 두 줄 연속으로 안되는 경우.
                    # elif (
                    #     prev_line     # 만약 line 별 x가 연속적으로 최대값 전 즈음이면 그냥 붙여준 다음 .붙여주고 glue # (보통 소주제 같은 것들)
                    #       and prev_line['x1'] < rightmost - self.__class__.NEWLINE_SPACE
                    #       and line_dict['x1'] < rightmost - self.__class__.NEWLINE_SPACE
                    #     ):
                    #     # if (not lines[-1][-1] in ".!?"  # 마지막 부분.
                    #     #     and not lines[-1][-2] in ".!?" # 마지막 부분 전 ("나 ' 부분)
                    #     #     ): 
                    #     #     lines[-1] = lines[-1]+"." # . added
                    #     #     no_add_period = True
                    #     sent = self.__class__.kiwi.glue(lines) 
                    #     file_wrote_flag = True
                    elif line_dict['x1'] < rightmost - self.__class__.NEWLINE_SPACE:
                        # if (not lines[-1][-1] in ".!?"  # 마지막 부분.
                        #     and not lines[-1][-2] in ".!?" # 마지막 부분 전 ("나 ' 부분)
                        #     ): 
                        #     lines[-1] = lines[-1]+"." # . added
                        #     no_add_period = True
                        sent = self.__class__.kiwi.glue(lines) 
                        file_wrote_flag = True
                    elif (line_dict['x1'] < rightmost - self.__class__.NEWLINE_SPACE # line에서 x가 맨 끝이 아닌데 .이 있으면 또 . 넣어주고 glue.
                          and line_text[-1] in ".!?" 
                        ):
                        sent = self.__class__.kiwi.glue(lines)
                        file_wrote_flag = True
                        # 만약 line에서 y가 큰 간격이라면 사이에 엔터 넣어주고. (y가 거의 끝까지가 아니거나 (이건 y 문제이긴 함... 이건 layout을 더 타이트하게 잡으면 됨.) )                    

                    if (len(page_lines)-1 == idx  # page's last line 
                        and line_dict['bottom'] < bottommost - self.__class__.PARAGRAPH_SPACE # 1. y 끝까지 오지 않은 경우. (이미 glue 되어 있음. 엔터)
                        ):
                        sent += "\n"
                        file_wrote_flag = True

                    # if (prev_line 
                    #     and not file_wrote_flag
                    #     and (line_dict['top'] - prev_line['bottom'] > self.__class__.LINE_HEIGHT_THRESHOLD)   # 2. 중간에 큰 공백이 삽입된 경우. (이미 glue 되어 있음. 엔터) 전꺼 확인 필요.
                    #     ):
                    #     lines[-2] = lines[-2]+ '\n'
                    # elif (prev_line 
                    #     and file_wrote_flag
                    #     and (line_dict['top'] - prev_line['bottom'] > self.__class__.LINE_HEIGHT_THRESHOLD)   # 2. 중간에 큰 공백이 삽입된 경우. (이미 glue 되어 있음. 엔터) 전꺼 확인 필요.
                    #     ):
                    #     sent += '\n'
                    if (prev_line 
                        and file_wrote_flag
                        and (line_dict['top'] - prev_line['bottom'] > self.__class__.LINE_HEIGHT_THRESHOLD)   # 2. 중간에 큰 공백이 삽입된 경우. (이미 glue 되어 있음. 엔터) 전꺼 확인 필요.
                        ):
                        sent = '\n' + sent
                    
                    if sent: # not ""
                        # add space if there is not a whitespace at the last of sent.
                        sent = sent + ' ' if (sent[-1] != " " and sent[-1] != "\n" ) else sent
                        ofile.write(sent)
                        lines.clear()   

                    
                    #     # if (
                    #     #     not no_add_period # period 없으면서.
                    #     #     and (not sent   # sent도 초기화 되어있고.
                    #     #     or (sent[-1] not in ".!?"  # 첫, 두번쨰 모두 이렇게 되어있고.
                    #     #         and sent[-2] not in ".!?")
                    #     #         )
                    #     #     ):    
                    #     #     ofile.write(".\n")  # 만약 소제목에 마침표가 없으면.
                    #     # else:
                    #     ofile.write("\n")  # 만약                 

                    prev_line = line_dict

                    #TODO: 만약 문단 변화가 있는 곳에서 엔터를 넣는 옵션을 준다면 거기서 엔터를 또 넣어주고.


                # for word in page.extract_words(use_text_flow=self.__class__.watermark,
                #                                x_tolerance=5,
                #                                y_tolerance=5,
                #                             #  split_at_punctutation=True,
                #                             extra_attrs=['size', 'stroking_color', 'non_stroking_color'],
                #                             ):
                    
                #     # print(f"word: {word}")

                #     # key = (round(word['size'],1), convert_fontcolor_to_tuple(word['non_stroking_color']))
                #     # if key not in font_prop:
                #     #     font_prop[key] = {'count':1, 'page_num':page.page_number, 'text':word['text']}
                #     # else:
                #     #     font_prop[key]['count'] += 1
                #     # new_dict = dict(sorted(font_prop.items(), key=lambda x: x[1]['count'], reverse=True))
                #     # for key, value in new_dict.items():
                #     #     print(key, value)

                #     x0, top, x1, bottom = map(lambda x : round(x, 2), (word['x0'], word['top'], word['x1'], word['bottom']))
                    
                #     cur_word = word['text']

                #     if double_quote_to_single_quote:
                #         cur_word = replace_to_single_double_quote(word['text'])                                
                    
                #     # 완전히 챕터가 종료되는 경우. 이건 나중에...
                #     if page_shifted is True:
                #         page_shifted = False
                #         if (re.search(r'[.!?]', self.last_word['text']) # .!?가 있었고, 
                #             and self.last_word['bottom'] < self.__class__.layout_location['bottommost'] - 55 # bottom에 놓는 것이 
                #             ):
                #             self.words_in_range.append('\n' + cur_word)
                #             self.update_last_word(word)
                #             continue
                        
                #         if (re.search(r'[.!?]', self.last_word['text']) 
                #             and self.last_word['x1'] < self.__class__.layout_location['rightmost'] - 10
                #             ):                               
                #             self.words_in_range.append('\n'+cur_word)
                #             self.update_last_word(word)
                #             continue
                        
                #         if self.spacing_condition(x0):
                #             self.words_in_range.pop()                               
                #             cur_word = spacing(self.words_in_range.pop() + cur_word)
                #             # append PAGE_ENDED in front of the spaced word.
                #             self.words_in_range.append(PAGE_ENDED)
                #             self.words_in_range.append(cur_word)
                #             self.update_last_word(word)
                #             continue
                    
                            
                #     if self.spacing_condition(x0) and (not self.over_line_height(top)):
                #         # No PAGE_ENEDED.
                #         cur_word = spacing(self.words_in_range.pop() + cur_word)

                #     if self.over_line_height(top): # 이것도 개인 설정.
                #         cur_word = '\n'+ cur_word
                    
                #     if need_newline(cur_word): 
                #         cur_word = cur_word+'\n'
                    
                #     self.words_in_range.append(cur_word)
                    
                #     # self.last_word update.
                #     self.update_last_word(word)
                
                # # 페이지 전환.
                # page_shifted = True
                    
                # # 워터마크 제거. 보완필요.
                # if self.__class__.watermark:
                #     self.eliminate_watermark()

                # self.words_in_range.append(PAGE_ENDED)
        
        print("pdf reading done.")
        # self.seperate_sentence() # 이거 녹여내면 속도 급상승할텐데...

    def write_reference_page(self):
        filename = self.__class__.title+ '_' + str(self.pages[0]) +'-' + str(self.pages[-1])+'_정보'+'.txt'
        with open(self.__class__.output_path + filename, 'w', encoding='utf-8') as f:
            f.write(f"책 제목: {self.__class__.title} \n")
            f.write(f"변환 페이지: {self.pages[0]}부터 {self.pages[-1]} \n")
            f.write("확인 요망 페이지:\n" + '\n'.join(self.image_page))
    

    def write_script_by_notepad(self):
        with open(self.__class__.output_path + self.__class__.title+ '_' + str(self.pages[0]) +'-' + str(self.pages[-1])+'.txt', 'w', encoding='utf-8') as f:
            f.write(' '.join(self.words_in_range))

    def write_script_by_docx(self):
        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = '맑은 고딕'
        font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

        page_count = -1
        p = document.add_paragraph('')
        for word in self.words_in_range:
            if word == PAGE_ENDED:
                page_count+=1
                if page_count in self.images.keys():
                    image_bytes = io.BytesIO()
                    self.images[page_count].save(image_bytes, format="PNG")
                    p.add_run().add_picture(image_bytes, width=Inches(5.0))
                continue
            p.add_run(word+' ')
        document.save(self.__class__.title+ '_' + str(self.pages[0]) +'-' + str(self.pages[-1])+'.docx')



if __name__ == "__main__":
    path = r"pdf\관계를 읽는 시간.pdf"        
    # path = r"pdf\sample.pdf"        
    pages = list(range(int(input('시작 페이지를 입력해주세요: ')), int(input('끝 페이지를 입력해주세요: '))+1))
    extractor = ScriptExtractor(pages)
    extractor.set_option(path, "C:/Users/82104/Desktop/", '관계를 읽는 시간 4차', 30,  excluded_lines_count=2) # 여기가 사용자가 입력해야 함.
    # extractor.set_layout()
    # extractor.read_pdf(double_quote_to_single_quote=True) # 이 부분도 사용자 입력 필요.
    extractor.set_new_layout()
    extractor.read_pdf_v2()
    # extractor.write_reference_page()
    # extractor.write_script_by_notepad()
    # extractor.write_script_by_docx()

# print(ScriptExtractor.path,
# ScriptExtractor.title,
# ScriptExtractor.watermark,
# ScriptExtractor.layout_pages,
# ScriptExtractor.categorize_by_font
# )


t2 = time.time()

print(f"소요시간: {t2-t1: .2f}초.")


               
        

# 네이버 백과사전 api로 하는 것은 외래어와 한자어 구분이 안될 수도 있기에 제외...

# from kiwipiepy import Kiwi

# kiwi = Kiwi(model_type='sbg')
# example_text = [
#                 # '이 책을 공동 집필한 우리 두 사람은 덴마크의 뉴스 프로그램 〈데드라인〉에 출연했다.',
#                 # '우리는 만나서 그동안 각자 봐온 것과 해온 일에 대해 오래 대화했고, 롤란드 파울센의 책, 『공허 노동, Empty Labor 』이 드디어 출판되었을 때 왜 더 큰 반향을 일으키지 않았을까 의아해했다.',
#                 # '그는 더 큰 정치적・경제적・사회적 미래 문제를 다루는 자리에 연사로 초청되었는데, 여기서 제시된 케인스의 전망은 오늘날 우리에게 라이트 못지않은 잘못된 예측으로 보인다.',
#                 # '(인생의 모든 단계에서) 보통 사람이 느리고 편한 속도에 맞춰 일하려는 경향이 있다는 데는 의문의 여지가 없다.',
#                 # "지금 우리 시대는 스탠리 큐브릭의 〈 2001 : 스페이스 오디세이〉 영화 속 설정 연도에서 20년이 훌쩍 지났지만, 감독의 상상이 실현되려면 아직 멀었다는 걸 생각하면 당황스러울 지경이다.",
#                 # "“예, 알겠습니다. 예….”"
#                 ]
# for i in example_text: # SSO (인용부호 열림), SSC (인용부호 닫힘)으로 바뀜.
#     return_text =  kiwi.tokenize(i)
#     print(f"return_text: {return_text}")


#TODO: sample 페이지로 layout 설정하는 과정 무조건 필요!
